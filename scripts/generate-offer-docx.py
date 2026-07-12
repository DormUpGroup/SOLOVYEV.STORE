"""Generate OFFER_SOLOVYEV_STORE.docx from structured content."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parents[1] / "docs" / "OFFER_SOLOVYEV_STORE.docx"
OUT_FALLBACK = Path(__file__).resolve().parents[1] / "docs" / "OFFER_SOLOVYEV_STORE_updated.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "1F2937") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
            for p in table.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Коммерческое предложение", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading("SOLOVYEV STORE — запуск и сопровождение сайта", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = [
        ("Исполнитель:", "Michael"),
        ("Заказчик:", "SOLOVYEV STORE"),
        ("Сайт:", "solovyev.store"),
        ("Дата:", "30 июня 2026"),
        ("Срок действия предложения:", "30 дней"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(value)

    doc.add_paragraph()

    # 1
    doc.add_heading("1. Предмет", level=2)
    doc.add_paragraph(
        "Запуск в продакшен обновлённого сайта SOLOVYEV STORE и сопровождение после запуска. "
        "Сайт уже разработан; работы включают настройку, наполнение, деплой и обучение."
    )
    p = doc.add_paragraph()
    p.add_run("Модель работы магазина: ").bold = True
    p.add_run("каталог онлайн → заказ через WhatsApp.")
    p = doc.add_paragraph()
    p.add_run("Не входит: ").bold = True
    p.add_run("онлайн-оплата, личные кабинеты покупателей.")

    # 2
    doc.add_heading("2. Пакет «Запуск» — €599–699 (разово)", level=2)
    doc.add_heading("2.1 Техническая часть", level=3)
    add_bullets(
        doc,
        [
            "Деплой на Vercel (продакшен)",
            "Подключение домена solovyev.store (настройка DNS)",
            "Настройка Supabase (база данных, хранилище фото)",
            "Перенос текущего каталога в админ-панель",
            "Подключение Google Analytics 4",
            "Смена пароля админки, базовая настройка безопасности",
            "Финальная проверка: корзина → WhatsApp, мобильная версия, SEO-страницы товаров",
        ],
    )
    doc.add_heading("2.2 Контент", level=3)
    add_bullets(
        doc,
        [
            "Загрузка до 30–50 товаров (фото, цена, размеры, бренд, статус) — в пакете «Стандарт»",
            "Настройка FAQ и страницы About (по материалам Заказчика)",
            "Настройка контактов, валюты (₪), объявлений в шапке сайта",
        ],
    )
    doc.add_heading("2.3 Обучение", level=3)
    doc.add_paragraph(
        "1 сессия до 60 минут: работа с админ-панелью /admin "
        "(добавление товара, фото, статусы SOLD / RESERVED / NEW DROP, кнопка Publish)."
    )
    doc.add_heading("2.4 Варианты цены", level=3)
    add_table(
        doc,
        ["Пакет", "Сумма", "Включает"],
        [
            ["Базовый", "€599 (~₪2 040)", "Готовая админ-панель без загрузки товаров в неё. Настраивать вручную."],
            ["Стандарт", "€699 (~₪2 400)", "Загрузка и настройка всех товаров + 2 раунда правок."],
        ],
    )
    doc.add_paragraph("Курс ориентир: 1 € ≈ 3,4 ₪.").italic = True

    # 3
    doc.add_heading("3. Пакет «Поддержка» — €100/мес", level=2)
    doc.add_heading("Входит ежемесячно", level=3)
    add_bullets(
        doc,
        [
            "Мониторинг работоспособности сайта",
            "До 2–3 мелких правок (текст, цена, контакт, баннер, мелкий баг)",
            "Консультации по админке (переписка / короткий созвон)",
            "Добавление до 5–10 товаров по материалам Заказчика",
        ],
    )
    doc.add_heading("Не входит (оплачивается отдельно)", level=3)
    add_bullets(
        doc,
        [
            "Крупные доработки и новые функции",
            "Массовая заливка каталога (50+ товаров)",
            "Реклама, SEO-продвижение, ведение Instagram",
            "Срочные задачи вне лимита — по отдельной договорённости",
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Минимальный срок: ").bold = True
    p.add_run("3 месяца (рекомендуется после запуска).")

    # 4
    doc.add_heading("4. Что нужно от Заказчика", level=2)
    add_table(
        doc,
        ["Материал", "Срок"],
        [
            ["Доступ к домену или данные регистратора", "до старта работ"],
            ["Подтверждение WhatsApp для заказов", "до старта работ"],
            ["Google Analytics ID (G-...)", "в течение 1 недели"],
            ["Товары: фото, цена, размеры, бренд", "в течение 1 недели"],
            ["Тексты FAQ и About (черновик допустим)", "в течение 1 недели"],
            ["До 60 мин на созвон для обучения", "после заливки каталога"],
        ],
    )
    doc.add_heading("Google Analytics ID — что это и зачем", level=3)
    p_ga1 = doc.add_paragraph()
    p_ga1.add_run("Что это: ").bold = True
    p_ga1.add_run(
        "бесплатный идентификатор счётчика Google Analytics 4 (GA4). Выглядит как G-XXXXXXXXXX (Measurement ID)."
    )
    p_ga2 = doc.add_paragraph()
    p_ga2.add_run("Зачем: ").bold = True
    p_ga2.add_run(
        "чтобы видеть статистику сайта после запуска — сколько человек зашло, какие товары смотрят, "
        "сколько раз добавили в корзину и перешли в WhatsApp. Без этого ID аналитика на сайте не подключается."
    )
    p_ga3 = doc.add_paragraph()
    p_ga3.add_run("Что будет видно: ").bold = True
    p_ga3.add_run(
        "просмотры страниц, добавления в корзину, переходы в WhatsApp, заявки Sell/Trade (в связке с настройкой сайта)."
    )
    p_ga4 = doc.add_paragraph()
    p_ga4.add_run("Кто создаёт: ").bold = True
    p_ga4.add_run(
        "рекомендуется Заказчик — сервис бесплатный, аккаунт остаётся у владельца магазина. "
        "По договорённости может создать Исполнитель с передачей доступа."
    )
    doc.add_paragraph().add_run("Как получить (кратко):").bold = True
    add_bullets(
        doc,
        [
            "Зайти на analytics.google.com под Google-аккаунтом Заказчика",
            "Создать property для solovyev.store → поток данных Web",
            "Скопировать Measurement ID (G-...) и отправить Исполнителю",
        ],
    )
    doc.add_paragraph("Задержка с материалами сдвигает срок запуска.")

    # 5
    doc.add_heading("5. Что получает Заказчик", level=2)
    add_bullets(
        doc,
        [
            "Сайт с каталогом (кроссовки / одежда / аксессуары), фильтрами и поиском",
            "SEO-страницы товаров — индексация в Google",
            "Корзина → структурированное сообщение в WhatsApp (размер, ссылка, номер заказа)",
            "Статусы SOLD / RESERVED / NEW DROP",
            "Страницы: Drops, Brands, About, FAQ, Sell/Trade, Privacy",
            "3 языка: иврит, русский, английский",
            "Админ-панель для самостоятельного управления каталогом",
            "Аналитика: просмотры, корзина, переходы в WhatsApp",
        ],
    )

    # 6
    doc.add_heading("6. Опциональные модули — launch-цена (фаза 2)", level=2)
    doc.add_paragraph(
        "Модули не обязательны для работы сайта. Каталог можно вести вручную через админ-панель."
    )
    p = doc.add_paragraph()
    p.add_run("Launch-цена ").bold = True
    p.add_run(
        "— скидка 50% при заказе в течение 30 календарных дней после запуска основного сайта (дата из п. 8). "
        "После истечения срока — стандартный прайс."
    )
    add_table(
        doc,
        ["Модуль", "Стандартный прайс", "Launch-цена (−50%)", "~₪"],
        [
            ["Instagram → сайт (автоимпорт постов)", "€600", "€300", "~₪1 020"],
            ["CRM: каталог + заказы + заявки (Admin + Google Sheets)", "€500", "€250", "~₪850"],
            ["Бандл: Instagram + CRM", "€800", "€400", "~₪1 360"],
            ["Кастомный checkout (PayPlus / Cardcom)", "€399", "€199", "~₪680"],
            ["ЛК покупателя (регистрация, профиль, заказы)", "€600", "€300", "~₪1 020"],
            ["ЛК сдатчика (consignment: статус вещей, выплаты)", "€900", "€450", "~₪1 530"],
            ["Бандл: ЛК покупателя + ЛК сдатчика", "€1 600", "€800", "~₪2 720"],
        ],
    )
    doc.add_heading("Instagram auto-import — что входит", level=3)
    add_bullets(
        doc,
        [
            "Подключение Instagram API (access token от Заказчика)",
            "Автоимпорт новых постов: фото, подпись → черновик товара в админке",
            "Настройка cron-синхронизации (по расписанию)",
            "Краткая инструкция по использованию",
        ],
    )
    doc.add_paragraph(
        "Подходит, если дропы публикуются в Instagram и нужно меньше ручной заливки на сайт."
    ).italic = True
    doc.add_heading("CRM: каталог + заказы + заявки — что входит", level=3)
    doc.add_paragraph(
        "Единый модуль: удобная админка + Google Таблица для учёта и резервной копии."
    )
    doc.add_paragraph().add_run("В админ-панели /admin").bold = True
    add_bullets(
        doc,
        [
            "Список заказов из WhatsApp checkout (фильтры, поиск, дата)",
            "Список заявок Sell/Trade с деталями и фото",
            "Статусы: новый / в работе / закрыт / отменён",
            "Карточка заказа / заявки — всё на одном экране",
            "Базовая статистика за период",
        ],
    )
    doc.add_paragraph().add_run("В Google Sheets (автосинхронизация)").bold = True
    add_bullets(
        doc,
        [
            "Лист «Каталог» — синк товаров с сайтом (название, цена, размеры, статус)",
            "Лист «Заказы» — WhatsApp checkout (дата, SS-*, товар, размер, цена, ссылка)",
            "Лист «Заявки» — Sell/Trade (контакт, описание, фото, статус)",
        ],
    )
    doc.add_paragraph("Краткая инструкция по работе с админкой и таблицей.")
    doc.add_paragraph(
        "Подходит для ежедневной работы в админке и параллельного учёта / отчётов в Google Sheets."
    ).italic = True
    doc.add_heading("Онлайн-оплата (PayPlus / Cardcom) — что входит", level=3)
    doc.add_paragraph(
        "Альтернатива Stripe для Израиля: приём карт в ₪, выплаты на израильский банк. "
        "Stripe напрямую для израильского бизнеса не подключается."
    )
    doc.add_paragraph().add_run("Кастомный checkout (€399 → €199 launch)").bold = True
    add_bullets(
        doc,
        [
            "Интеграция PayPlus или Cardcom",
            "Корзина → оплата картой → страницы success / cancel",
            "Тестовый режим + запуск на проде",
            "Оформление оплаты в UI сайта (API), без «чужой» страницы",
            "Кнопки и flow в стиле SOLOVYEV STORE",
            "Webhook от платёжки, статус оплаты в админке",
            "Лог заказов с оплатой (в связке с CRM-модулем — по желанию)",
            "Bit / Apple Pay — если включены в тарифе провайдера",
        ],
    )
    doc.add_paragraph().add_run("Не входит (на стороне Заказчика):").bold = True
    add_bullets(
        doc,
        [
            "Договор с PayPlus / Cardcom / Tranzila, merchant account, terminal",
            "עוסק מורשה, банковский счёт, комиссии шлюза (~1.5–3.5% + абонплата)",
            "חשבונית מס / интеграция с бухгалтерией",
        ],
    )
    doc.add_paragraph(
        "Рекомендуется после запуска сайта с WhatsApp checkout."
    ).italic = True
    doc.add_heading("Личные кабинеты — что входит", level=3)
    doc.add_paragraph(
        "Кабинет владельца уже входит в пакет «Запуск» — это админ-панель /admin (каталог, настройки, FAQ). "
        "Отдельно не оплачивается."
    )
    doc.add_paragraph().add_run("ЛК покупателя (€600 → €300 launch)").bold = True
    add_bullets(
        doc,
        [
            "Регистрация и вход (email)",
            "Профиль: имя, телефон, адрес доставки",
            "История заказов (из CRM / WhatsApp checkout; при оплате — из платёжки)",
            "Избранное (по согласованию)",
        ],
    )
    doc.add_paragraph().add_run("ЛК сдатчика / consignment (€900 → €450 launch)").bold = True
    add_bullets(
        doc,
        [
            "Личный кабинет для тех, кто сдаёт вещи на продажу",
            "Список вещей, статусы (на продаже / продано / выплачено)",
            "Сумма к выплате, история сделок",
            "Связка с Sell/Trade и админкой владельца",
        ],
    )
    doc.add_paragraph().add_run("Бандл: ЛК покупателя + ЛК сдатчика (€1 600 → €800 launch)").bold = True
    doc.add_paragraph("— оба модуля со скидкой.")
    doc.add_paragraph(
        "Рекомендуется после запуска основного сайта. Без CRM или онлайн-оплаты часть данных в ЛК "
        "заполняется вручную или из WhatsApp."
    ).italic = True
    doc.add_heading("Сроки воплощения модулей", level=3)
    add_table(
        doc,
        ["Модуль / услуга", "Срок", "Отсчёт с"],
        [
            ["Instagram → сайт", "5–7 раб. дней", "предоплата 30% + Instagram access token"],
            ["CRM (Admin + Google Sheets)", "7–10 раб. дней", "предоплата 30% + доступ к Google Таблице"],
            ["Кастомный checkout (PayPlus / Cardcom)", "7–14 раб. дней", "предоплата 30% + live API keys от провайдера"],
            ["ЛК покупателя", "14–21 раб. день", "предоплата 30% + согласованное ТЗ"],
            ["ЛК сдатчика (consignment)", "21–28 раб. дней", "предоплата 30% + согласованное ТЗ"],
            ["Бандл: ЛК покупателя + ЛК сдатчика", "21–30 раб. дней", "предоплата 30% + ТЗ по обоим модулям"],
            ["Бандл: Instagram + CRM", "10–14 раб. дней", "предоплата 30% + все доступы по модулям"],
            ["Логотип", "3–5 раб. дней", "бриф / референсы от Заказчика"],
            ["Блок дизайна", "2–4 раб. дня", "согласованное ТЗ"],
            ["Заливка товаров (до 10 шт.)", "1–2 раб. дня", "фото и данные от Заказчика"],
        ],
    )
    doc.add_paragraph(
        "Сроки — рабочие дни. Задержка доступов или материалов от Заказчика сдвигает срок пропорционально."
    ).italic = True
    doc.add_heading("Условия заказа модулей", level=3)
    add_bullets(
        doc,
        [
            "Заказ по launch-цене — до истечения 30-дневного срока",
            "Оплата модулей: 30% при заказе, 40% после правок, 30% после сдачи и проверки",
            "Сроки воплощения — по таблице выше; старт после предоплаты и получения доступов",
            "Модули можно заказать по отдельности или бандлом",
        ],
    )
    doc.add_heading("Прочие дополнительные услуги (вне launch-акции)", level=3)
    add_table(
        doc,
        ["Услуга", "Цена"],
        [
            ["Заливка товаров сверх лимита пакета «Запуск»", "€15 / товар"],
            ["Создание/смена логотипа (2–3 варианта, 2 раунда правок, PNG + SVG)", "€39"],
            ["Новый блок / доработка дизайна", "€79 / блок"],
        ],
    )

    # 7-11
    doc.add_heading("7. Расходы Заказчика (не входят в оплату Исполнителю)", level=2)
    add_table(
        doc,
        ["Статья", "Стоимость"],
        [
            ["Домен .store", "у регистратора"],
            ["Vercel + Supabase (старт)", "€0 (бесплатный тариф)"],
        ],
    )

    doc.add_heading("8. Сроки", level=2)
    add_table(
        doc,
        ["Этап", "Срок"],
        [
            ["Старт после предоплаты и получения доступов", "1–2 дня"],
            ["Запуск в продакшен", "7–14 рабочих дней после получения всех материалов"],
            ["Обучение", "в течение 3 дней после запуска"],
        ],
    )

    doc.add_heading("9. Оплата", level=2)
    add_table(
        doc,
        ["Этап", "Сумма", "Когда"],
        [
            ["Предоплата (запуск)", "30%", "перед началом работ"],
            ["Промежуточный платёж", "40%", "после заливки каталога и готовности staging-версии"],
            ["Финальный платёж", "30%", "в день запуска / после приёмки"],
            ["Поддержка", "€100/мес", "1-го числа, постоплата или предоплата за месяц"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Способы оплаты: ").bold = True
    p.add_run("банковский перевод (SEPA / IBAN) / PayPal.")

    doc.add_heading("10. Приёмка", level=2)
    doc.add_paragraph(
        "Сайт считается принятым, если в течение 5 рабочих дней после запуска Заказчик не направил "
        "письменный список критичных замечаний. Некритичные правки (тексты, фото) — в рамках пакета или поддержки."
    )

    doc.add_heading("11. Подписи", level=2)
    doc.add_paragraph().add_run("Исполнитель").bold = True
    doc.add_paragraph("Имя: Michael _____________________")
    doc.add_paragraph("Подпись: _____________________")
    doc.add_paragraph("Дата: ________________________")
    doc.add_paragraph()
    doc.add_paragraph().add_run("Заказчик (SOLOVYEV STORE)").bold = True
    doc.add_paragraph("Имя: Gosha ___________________")
    doc.add_paragraph("Подпись: _____________________")
    doc.add_paragraph("Дата: ________________________")
    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Предложение не является публичной офертой до подписания обеими сторонами."
    )
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT)
        print(f"Saved: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"Saved (original file locked): {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
