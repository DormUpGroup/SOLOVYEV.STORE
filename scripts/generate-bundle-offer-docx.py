"""Generate OFFER_BUSINESS_BUNDLE.docx — phase 2 modules package."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parents[1] / "docs" / "OFFER_BUSINESS_BUNDLE.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "1F2937") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
            for p in table.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Коммерческое предложение", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_heading("SOLOVYEV STORE — пакет «Бизнес Pro» (фаза 2)", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = [
        ("Исполнитель:", "Michael"),
        ("Заказчик:", "SOLOVYEV STORE"),
        ("Сайт:", "solovyev.store"),
        ("Дата:", "30 июня 2026"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(value)

    doc.add_paragraph()

    doc.add_heading("1. Состав пакета", level=2)
    add_bullets(
        doc,
        [
            "Личный кабинет для покупателей (регистрация, профиль, история заказов)",
            "CRM в админке: заказы, заявки Sell/Trade, синхронизация с Google Sheets",
            "CRM маржа: цена закупа, цена продажи, прибыль по позиции, дашборд «общий лут»",
            "Автоматизация Instagram → сайт (посты в каталог)",
        ]
    )

    doc.add_paragraph(
        "Кабинет владельца (админка /admin для каталога и настроек) входит в основной пакет запуска "
        "€599–699 и в этот пакет отдельно не входит."
    )

    doc.add_heading("2. Модули по отдельности", level=2)
    add_table(
        doc,
        ["Модуль", "Launch (−50%)", "Стандарт", "~₪ launch"],
        [
            ["Instagram → сайт", "€300", "€600", "~₪1 020"],
            ["ЛК покупателя", "€300", "€600", "~₪1 020"],
            ["CRM (заказы, заявки, Sheets)", "€250", "€500", "~₪850"],
            ["CRM маржа (закуп / продажа / прибыль / дашборд)", "€300", "€600", "~₪1 020"],
        ],
    )

    p_sum = doc.add_paragraph()
    p_sum.add_run("Сумма по отдельности: ").bold = True
    p_sum.add_run("€1 150 launch / €2 300 стандарт.")

    doc.add_paragraph(
        "CRM маржа — блок поверх базового CRM: поля «купили за / продали за», маржа %, "
        "прибыль по каждой позиции, итог за период в админке."
    )

    doc.add_heading("3. Бандл «под ключ» (рекомендуется)", level=2)
    add_table(
        doc,
        ["Пакет", "Launch", "Стандарт"],
        [
            ["Всё вместе (Instagram + ЛК + CRM + маржа)", "€950", "€1 900"],
        ],
    )
    doc.add_paragraph("Экономия ~€200 против поштучной покупки модулей.").italic = True

    doc.add_heading("4. Сроки и оплата", level=2)
    add_bullets(
        doc,
        [
            "Срок воплощения: 5–6 недель (часть работ параллельно)",
            "Старт после предоплаты 30% и получения доступов: Instagram token, ТЗ по марже, согласование ЛК",
            "Оплата: 30% при заказе → 40% после промежуточной сдачи → 30% после приёмки",
            "Launch-цена действует 30 календарных дней после запуска основного сайта",
        ]
    )

    doc.add_paragraph("Курс ориентир: 1 € ≈ 3,4 ₪.").italic = True

    doc.add_heading("5. Подписи", level=2)
    doc.add_paragraph().add_run("Исполнитель").bold = True
    doc.add_paragraph("Имя: Michael _____________________")
    doc.add_paragraph("Подпись: _____________________")
    doc.add_paragraph("Дата: ________________________")
    doc.add_paragraph()
    doc.add_paragraph().add_run("Заказчик (SOLOVYEV STORE)").bold = True
    doc.add_paragraph("Имя: _________________________")
    doc.add_paragraph("Подпись: _____________________")
    doc.add_paragraph("Дата: ________________________")

    foot = doc.add_paragraph(
        "Дополнение к основному коммерческому предложению SOLOVYEV STORE."
    )
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
